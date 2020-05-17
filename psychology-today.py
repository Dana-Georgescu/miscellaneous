import requests, bs4, os, re, pprint, pyperclip
from docx import Document

url_blog = 'https://www.psychologytoday.com/us/blog/in-practice-0?page='
url_root = 'https://www.psychologytoday.com'
#! python3.7
‘’’ psychologie-today.py - Downloads all articles of this blog «https://www.psychologytoday.com/us/blog/in-practice» to a docx document

headers = {'User-Agent': "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_12_6) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/79.0.3945.117 Safari/537.36"}
inpractice = 'https://www.psychologytoday.com/us/blog/in-practice-0?page='

document = Document('articole_in_practice6.docx')
document.add_heading('Articole In Practice', 0)

def get_all_article_urls(url):
    page = 0
    url_list = []
    article_title = []
    for page in range(10):
        res = requests.get(url+str(page), headers=headers)
        res.raise_for_status()
        soup = bs4.BeautifulSoup(res.text, 'html.parser')
        tags = soup.select('h2.blog_entry__title a')
        for tag in tags:
            article_title.append(tag.getText())
        for tag in tags:
            article = tag.get('href')
            url_list.append(url_root + article)
    return url_list



def copy_text(url):
    #url = 'https://www.psychologytoday.com/us/blog/in-practice/202001/how-get-emotional-support-when-you-feel-you-have-none'
    res = requests.get(url, headers=headers)
    soup = bs4.BeautifulSoup(res.text, 'html.parser')
    title = soup.select('.blog_entry--full__title')[0].getText()
    subtitle = soup.select('.blog_entry--full__subtitle')[0].getText().strip()
    text = soup.select('div .field.field-name-body.field-type-text-with-summary.field-label-hidden p')
    document.add_heading(title, level=1)
    document.add_heading(subtitle, level=2)
    for i in range(len(text)):
        document.add_paragraph(text[i].getText())

urls = get_all_article_urls(inpractice)
print('number of articles ', len(urls))

for url in urls[281:290]:
    print(url)
    copy_text(url)
    
    
'''p = document.add_paragraph(text[0])
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True'''
document.save('articole_in_practice7.docx')
